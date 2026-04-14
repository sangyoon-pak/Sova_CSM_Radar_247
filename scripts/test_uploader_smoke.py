#!/usr/bin/env python3
"""Smoke test for document uploader with txt/docx/pdf fixtures."""

from __future__ import annotations

import json
import tempfile
import urllib.request
import uuid
from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas


BASE_URL = "http://127.0.0.1:8000"


def _multipart_body(field_name: str, filename: str, content: bytes, boundary: str) -> bytes:
    lines: list[bytes] = []
    lines.append(f"--{boundary}\r\n".encode())
    lines.append(
        (
            f'Content-Disposition: form-data; name="{field_name}"; filename="{filename}"\r\n'
            "Content-Type: application/octet-stream\r\n\r\n"
        ).encode()
    )
    lines.append(content)
    lines.append(b"\r\n")
    lines.append(f"--{boundary}--\r\n".encode())
    return b"".join(lines)


def _upload(path: Path) -> dict:
    boundary = f"----WebKitFormBoundary{uuid.uuid4().hex}"
    body = _multipart_body("file", path.name, path.read_bytes(), boundary)
    req = urllib.request.Request(
        f"{BASE_URL}/docs/upload",
        data=body,
        method="POST",
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
    )
    with urllib.request.urlopen(req, timeout=30) as resp:
        raw = resp.read().decode("utf-8")
    return json.loads(raw)


def _fetch_docs(limit: int = 30) -> list[dict]:
    with urllib.request.urlopen(f"{BASE_URL}/kb/documents?limit={limit}", timeout=30) as resp:
        raw = resp.read().decode("utf-8")
    data = json.loads(raw)
    return data if isinstance(data, list) else []


def _assert(condition: bool, msg: str) -> None:
    if not condition:
        raise AssertionError(msg)


def _make_txt(path: Path) -> None:
    path.write_text(
        "TXT Upload Test\n"
        "Customer asks about API push campaign CID mismatch and troubleshooting.\n",
        encoding="utf-8",
    )


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("DOCX Upload Test", 0)
    doc.add_paragraph("Client asks about webhook delay and campaign attribute sync errors.")
    doc.add_paragraph("Need grounded answer with references.")
    doc.save(str(path))


def _make_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path))
    c.drawString(72, 800, "PDF Upload Test")
    c.drawString(72, 780, "Customer asks about SDK integration issue and callback behavior.")
    c.drawString(72, 760, "Please provide troubleshooting path and references.")
    c.save()


def main() -> None:
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        txt = root / "uploader_smoke_sample.txt"
        docx = root / "uploader_smoke_sample.docx"
        pdf = root / "uploader_smoke_sample.pdf"
        _make_txt(txt)
        _make_docx(docx)
        _make_pdf(pdf)

        results = [_upload(txt), _upload(docx), _upload(pdf)]
        for res in results:
            _assert(bool(res.get("saved")), f"upload failed: {res}")
            _assert(bool(res.get("filename")), f"missing filename: {res}")
            _assert(int(res.get("chars") or 0) > 20, f"extracted text too short: {res}")

        names = [str(r.get("filename")) for r in results]
        _assert(len(set(names)) == 3, f"filenames collided: {names}")

        docs = _fetch_docs(limit=50)
        by_id = {int(d.get("id")): d for d in docs if d.get("id") is not None}
        for res in results:
            did = int(res.get("doc_id") or 0)
            _assert(did in by_id, f"doc id not found in registry: {did}")
            row = by_id[did]
            p = str(row.get("path") or "")
            _assert(p.endswith(".md"), f"normalized path should be .md: {p}")
            _assert("uploader_smoke_sample" in p, f"unexpected output path: {p}")

        print("Uploader smoke test passed: txt/docx/pdf upload + registry checks OK.")


if __name__ == "__main__":
    main()
